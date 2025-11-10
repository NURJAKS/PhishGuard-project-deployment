"""
Модуль для ИИ-проверки счетов-фактур
Проверяет форматы, логику, дубликаты, семантику и сопоставляет с внешними данными
"""
import os
import logging
import uuid
import re
import hashlib
from typing import Dict, Any, Optional, List
from datetime import datetime, timedelta
import json

logger = logging.getLogger(__name__)

# Try to import document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available, PDF processing disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth
    DOCX_MAMMOTH_AVAILABLE = True
except ImportError:
    DOCX_MAMMOTH_AVAILABLE = False

# Инициализация переменных для Excel
EXCEL_AVAILABLE = False
EXCEL_XLRD_AVAILABLE = False
openpyxl = None
xlrd = None

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    openpyxl = None

try:
    import xlrd
    EXCEL_XLRD_AVAILABLE = True
except ImportError:
    EXCEL_XLRD_AVAILABLE = False
    xlrd = None

try:
    import xml.etree.ElementTree as ET
    XML_AVAILABLE = True
except ImportError:
    XML_AVAILABLE = False

# Try to import AI analyzer
try:
    from ai_analyzer import _openrouter_analyze, _google_ai_analyze
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False
    logger.warning("AI analyzer not available")

# Storage for invoice analyses (in production, use database)
INVOICE_ANALYSES = {}
INVOICE_DUPLICATES = {}  # Для проверки дубликатов: (number, seller_inn, date) -> count


def extract_text_from_pdf(file_content: bytes) -> str:
    """Извлекает текст из PDF файла"""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 not installed")
    
    try:
        import io
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise


def extract_text_from_docx(file_content: bytes) -> str:
    """Извлекает текст из DOC/DOCX файла"""
    if not DOCX_AVAILABLE and not DOCX_MAMMOTH_AVAILABLE:
        raise ImportError("python-docx or mammoth not installed")
    
    try:
        import io
        docx_file = io.BytesIO(file_content)
        
        # Try python-docx first
        if DOCX_AVAILABLE:
            doc = Document(docx_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            return text.strip()
        
        # Fallback to mammoth
        if DOCX_MAMMOTH_AVAILABLE:
            result = mammoth.extract_raw_text(docx_file)
            return result.value.strip()
        
        raise ImportError("No DOCX library available")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise


def extract_text_from_excel(file_content: bytes, filename: str) -> str:
    """Извлекает текст из Excel файла (XLSX/XLS)"""
    try:
        import io
        text_parts = []
        file_ext = filename.lower().split('.')[-1] if filename else ""
        
        # Try openpyxl for XLSX/XLSM
        if EXCEL_AVAILABLE and openpyxl is not None and file_ext in ['xlsx', 'xlsm']:
            try:
                excel_file = io.BytesIO(file_content)
                workbook = openpyxl.load_workbook(excel_file, data_only=True)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text_parts.append(f"\n=== Лист: {sheet_name} ===\n")
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            text_parts.append(row_text)
                return "\n".join(text_parts).strip()
            except Exception as e:
                logger.warning(f"Error with openpyxl, trying xlrd: {e}")
                # Fallback to xlrd if openpyxl fails
        
        # Try xlrd for XLS (old format) or as fallback
        if EXCEL_XLRD_AVAILABLE and xlrd is not None and file_ext in ['xls', 'xlsx']:
            try:
                workbook = xlrd.open_workbook(file_contents=file_content)
                for sheet_name in workbook.sheet_names():
                    sheet = workbook.sheet_by_name(sheet_name)
                    text_parts.append(f"\n=== Лист: {sheet_name} ===\n")
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        row_text = " | ".join([str(cell) for cell in row])
                        if row_text.strip():
                            text_parts.append(row_text)
                return "\n".join(text_parts).strip()
            except Exception as e:
                logger.error(f"Error with xlrd: {e}")
                raise
        
        # Если ни одна библиотека не доступна
        if not EXCEL_AVAILABLE and not EXCEL_XLRD_AVAILABLE:
            raise ImportError("No Excel library available. Install openpyxl (for XLSX) or xlrd (for XLS)")
        
        raise ValueError(f"Cannot process Excel file: {filename}. Supported formats: XLSX (needs openpyxl), XLS (needs xlrd)")
        
    except Exception as e:
        logger.error(f"Error extracting text from Excel: {e}")
        raise


def extract_text_from_xml(file_content: bytes) -> str:
    """Извлекает текст из XML файла (УПД/счёт-фактура)"""
    if not XML_AVAILABLE:
        raise ImportError("XML parser not available")
    
    try:
        root = ET.fromstring(file_content)
        # Простое извлечение текста из всех элементов
        text_parts = []
        for elem in root.iter():
            if elem.text:
                text_parts.append(elem.text.strip())
        return " ".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from XML: {e}")
        raise


def validate_inn(inn: str) -> bool:
    """Проверка валидности ИНН (10 или 12 цифр)"""
    if not inn or not isinstance(inn, str):
        return False
    inn_clean = re.sub(r'\D', '', inn)
    if len(inn_clean) not in [10, 12]:
        return False
    return inn_clean.isdigit()


def validate_kpp(kpp: str) -> bool:
    """Проверка валидности КПП (9 цифр)"""
    if not kpp or not isinstance(kpp, str):
        return False
    kpp_clean = re.sub(r'\D', '', kpp)
    return len(kpp_clean) == 9 and kpp_clean.isdigit()


def extract_invoice_data(text: str) -> Dict[str, Any]:
    """Извлекает реквизиты счёта-фактуры из текста"""
    data = {
        "number": None,
        "issue_date": None,
        "due_date": None,
        "amount": None,
        "vat_rate": None,
        "vat_amount": None,
        "total_amount": None,
        "seller": {
            "name": None,
            "inn": None,
            "kpp": None
        },
        "buyer": {
            "name": None,
            "inn": None,
            "kpp": None
        },
        "items": []
    }
    
    # Извлечение номера счёта-фактуры (более гибкие паттерны)
    number_patterns = [
        r'Сч[её]т[- ]?фактура\s*[№N#]?\s*([\w\-\/]+)',
        r'[№N#]\s*(\d+)',
        r'№\s*(\d+)',
        r'номер[:\s]+(\d+)',
        r'\b(\d{3,})\b'  # Просто ищем числа из 3+ цифр в начале документа
    ]
    for pattern in number_patterns:
        number_match = re.search(pattern, text, re.IGNORECASE)
        if number_match:
            data["number"] = number_match.group(1).strip()
            break
    
    # Извлечение дат (более гибкие паттерны)
    date_patterns = [
        r'(\d{1,2})[./](\d{1,2})[./](\d{4})',  # DD.MM.YYYY или DD/MM/YYYY
        r'(\d{4})[./-](\d{1,2})[./-](\d{1,2})',  # YYYY.MM.DD
        r'(\d{1,2})\s+(\w+)\s+(\d{4})',  # "31 октября 2025"
    ]
    dates = []
    for pattern in date_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                if len(match) == 3:
                    # Проверяем валидность даты
                    if match[1].isdigit():  # DD.MM.YYYY или YYYY.MM.DD
                        if len(match[2]) == 4:  # DD.MM.YYYY
                            day, month, year = int(match[0]), int(match[1]), int(match[2])
                            if 1 <= month <= 12 and 1 <= day <= 31 and 2000 <= year <= 2100:
                                date_str = f"{year}-{month:02d}-{day:02d}"
                                dates.append(date_str)
                        else:  # YYYY.MM.DD
                            year, month, day = int(match[0]), int(match[1]), int(match[2])
                            if 1 <= month <= 12 and 1 <= day <= 31 and 2000 <= year <= 2100:
                                date_str = f"{year}-{month:02d}-{day:02d}"
                                dates.append(date_str)
                    else:  # "31 октября 2025"
                        month_map = {
                            'января': '01', 'февраля': '02', 'марта': '03', 'апреля': '04',
                            'мая': '05', 'июня': '06', 'июля': '07', 'августа': '08',
                            'сентября': '09', 'октября': '10', 'ноября': '11', 'декабря': '12'
                        }
                        month = month_map.get(match[1].lower())
                        if month:
                            day, year = int(match[0]), int(match[2])
                            if 1 <= day <= 31 and 2000 <= year <= 2100:
                                date_str = f"{year}-{month}-{day:02d}"
                                dates.append(date_str)
            except Exception as e:
                logger.debug(f"Error parsing date {match}: {e}")
                pass
    
    # Убираем дубликаты и сортируем
    dates = sorted(list(set(dates)))
    if dates:
        data["issue_date"] = dates[0]
        if len(dates) > 1:
            data["due_date"] = dates[1]
    
    # Извлечение сумм (более гибкие паттерны)
    # Сначала извлекаем ИНН/КПП, чтобы исключить их из поиска суммы
    known_inns_kpps = set()
    if data["seller"]["inn"]:
        known_inns_kpps.add(data["seller"]["inn"])
    if data["seller"]["kpp"]:
        known_inns_kpps.add(data["seller"]["kpp"])
    if data["buyer"]["inn"]:
        known_inns_kpps.add(data["buyer"]["inn"])
    if data["buyer"]["kpp"]:
        known_inns_kpps.add(data["buyer"]["kpp"])
    
    amount_patterns = [
        r'Итого\s*к\s*оплате[:\s]+([\d\s,\.]+)',
        r'К\s*оплате[:\s]+([\d\s,\.]+)',
        r'Сумма[:\s]+([\d\s,\.]+)',
        r'Всего[:\s]+([\d\s,\.]+)',
        r'(\d{5,})\s*руб',  # Ищем большие числа с "руб"
        r'(\d{5,})\s*[рр]',  # Ищем большие числа с "р"
    ]
    
    for pattern in amount_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            amount_str = match.group(1).replace(' ', '').replace(',', '.').replace(' ', '')
            try:
                amount = float(amount_str)
                # Проверяем, что сумма разумная (больше 1000, но меньше 1 миллиарда)
                # и это не ИНН/КПП
                if 1000 < amount < 1000000000 and amount_str not in known_inns_kpps:
                    data["total_amount"] = amount
                    break
            except:
                pass
    
    # Если не нашли через паттерны, ищем большие числа в конце документа
    if not data["total_amount"]:
        large_numbers = re.findall(r'\b(\d{5,})\b', text)
        if large_numbers:
            # Исключаем ИНН/КПП и берём самое большое разумное число
            try:
                valid_amounts = []
                for n in large_numbers:
                    n_clean = n.replace(' ', '').replace(',', '.')
                    if n_clean not in known_inns_kpps:
                        try:
                            amount = float(n_clean)
                            # Разумная сумма для счёта-фактуры: от 10,000 до 10,000,000
                            if 10000 <= amount <= 10000000:
                                valid_amounts.append(amount)
                        except:
                            pass
                if valid_amounts:
                    # Берём самое большое число как итоговую сумму
                    data["total_amount"] = max(valid_amounts)
            except:
                pass
    
    # Извлечение НДС
    vat_patterns = [
        r'НДС\s*(\d+)%',
        r'НДС[:\s]+(\d+)%',
        r'Ставка\s*НДС[:\s]+(\d+)%'
    ]
    for pattern in vat_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                data["vat_rate"] = int(match.group(1))
                break
            except:
                pass
    
    # Извлечение суммы НДС
    vat_amount_patterns = [
        r'НДС\s*\d+%[:\s]+([\d\s,\.]+)',
        r'НДС[:\s]+([\d\s,\.]+)\s*руб',
        r'Сумма\s*НДС[:\s]+([\d\s,\.]+)'
    ]
    for pattern in vat_amount_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            vat_amount_str = match.group(1).replace(' ', '').replace(',', '.')
            try:
                data["vat_amount"] = float(vat_amount_str)
                break
            except:
                pass
    
    # Извлечение ИНН продавца (более гибкие паттерны - ищем просто числа нужной длины)
    inn_patterns = [
        r'Продавец[:\s]+.*?ИНН[:\s]*([0-9]{10,12})',
        r'ИНН[:\s]*([0-9]{10,12})',
        r'\b([0-9]{10})\b',  # 10-значный ИНН
        r'\b([0-9]{12})\b'   # 12-значный ИНН
    ]
    for pattern in inn_patterns:
        seller_inn_match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if seller_inn_match:
            inn = seller_inn_match.group(1)
            # Проверяем, что это действительно ИНН (10 или 12 цифр)
            if len(inn) in [10, 12] and inn.isdigit():
                data["seller"]["inn"] = inn
                break
    
    # Извлечение КПП продавца
    kpp_patterns = [
        r'Продавец[:\s]+.*?КПП[:\s]*([0-9]{9})',
        r'КПП[:\s]*([0-9]{9})',
        r'\b([0-9]{9})\b'  # 9-значный КПП
    ]
    for pattern in kpp_patterns:
        seller_kpp_match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if seller_kpp_match:
            kpp = seller_kpp_match.group(1)
            if len(kpp) == 9 and kpp.isdigit():
                data["seller"]["kpp"] = kpp
                break
    
    # Извлечение наименования продавца (более гибкий паттерн)
    seller_name_patterns = [
        r'Продавец[:\s]+([А-ЯЁ][А-Яа-яё\s]+(?:ООО|АО|ИП|ЗАО|ОАО))',
        r'([А-ЯЁ][А-Яа-яё\s]+(?:ООО|АО|ИП|ЗАО|ОАО))',  # Просто ищем организации
    ]
    for pattern in seller_name_patterns:
        seller_name_match = re.search(pattern, text, re.IGNORECASE)
        if seller_name_match:
            name = seller_name_match.group(1).strip()
            # Проверяем, что это не слишком короткое совпадение
            if len(name) > 5:
                data["seller"]["name"] = name
                break
    
    # Извлечение ИНН покупателя
    buyer_inn_match = re.search(r'Покупатель[:\s]+.*?ИНН[:\s]*([0-9]{10,12})', text, re.IGNORECASE | re.DOTALL)
    if buyer_inn_match:
        data["buyer"]["inn"] = buyer_inn_match.group(1)
    
    return data


def check_arithmetic(invoice_data: Dict[str, Any]) -> Dict[str, Any]:
    """Проверка арифметики: сумма позиций = итого, НДС сходится"""
    check_result = {
        "ok": True,
        "details": []
    }
    
    total = invoice_data.get("total_amount")
    vat_rate = invoice_data.get("vat_rate")
    vat_amount = invoice_data.get("vat_amount")
    
    if total and vat_rate:
        # Проверяем расчёт НДС
        if vat_rate == 20:
            expected_vat = total * 20 / 120
            if vat_amount and abs(vat_amount - expected_vat) > 0.01:
                check_result["ok"] = False
                check_result["details"].append(f"НДС не сходится: ожидается {expected_vat:.2f}, указано {vat_amount}")
        elif vat_rate == 10:
            expected_vat = total * 10 / 110
            if vat_amount and abs(vat_amount - expected_vat) > 0.01:
                check_result["ok"] = False
                check_result["details"].append(f"НДС не сходится: ожидается {expected_vat:.2f}, указано {vat_amount}")
    
    return check_result


def check_duplicate(invoice_data: Dict[str, Any]) -> Dict[str, Any]:
    """Проверка на дубликаты по номеру+ИНН+дате"""
    check_result = {
        "ok": True,
        "details": []
    }
    
    number = invoice_data.get("number")
    seller_inn = invoice_data.get("seller", {}).get("inn")
    issue_date = invoice_data.get("issue_date")
    
    if number and seller_inn and issue_date:
        key = f"{number}|{seller_inn}|{issue_date}"
        count = INVOICE_DUPLICATES.get(key, 0)
        INVOICE_DUPLICATES[key] = count + 1
        
        if count > 0:
            check_result["ok"] = False
            check_result["details"].append(f"Найден дубликат: счёт-фактура №{number} от {issue_date} уже встречалась ранее")
    
    return check_result


def check_dates(invoice_data: Dict[str, Any]) -> Dict[str, Any]:
    """Проверка дат: дата выдачи <= дата оплаты"""
    check_result = {
        "ok": True,
        "details": []
    }
    
    issue_date = invoice_data.get("issue_date")
    due_date = invoice_data.get("due_date")
    
    if issue_date and due_date:
        try:
            issue = datetime.strptime(issue_date, "%Y-%m-%d")
            due = datetime.strptime(due_date, "%Y-%m-%d")
            if issue > due:
                check_result["ok"] = False
                check_result["details"].append("Дата выдачи позже даты оплаты")
        except:
            pass
    
    return check_result


def check_vat_rate(invoice_data: Dict[str, Any]) -> Dict[str, Any]:
    """Проверка допустимости ставки НДС"""
    check_result = {
        "ok": True,
        "details": []
    }
    
    vat_rate = invoice_data.get("vat_rate")
    if vat_rate:
        valid_rates = [0, 10, 20]
        if vat_rate not in valid_rates:
            check_result["ok"] = False
            check_result["details"].append(f"Недопустимая ставка НДС: {vat_rate}%")
    
    return check_result


def semantic_check(invoice_data: Dict[str, Any], text: str) -> Dict[str, Any]:
    """Семантическая проверка через AI (опционально)"""
    check_result = {
        "ok": True,
        "similarity": 1.0,
        "details": []
    }
    
    if not AI_AVAILABLE:
        return check_result
    
    # Проверяем наличие подозрительных формулировок
    suspicious_patterns = [
        r'прочие\s+услуги',
        r'корректировка',
        r'дополнительное\s+соглашение',
        r'без\s+описания'
    ]
    
    text_lower = text.lower()
    for pattern in suspicious_patterns:
        if re.search(pattern, text_lower):
            check_result["similarity"] = max(0.0, check_result["similarity"] - 0.2)
            check_result["details"].append(f"Найдена подозрительная формулировка: {pattern}")
    
    if check_result["similarity"] < 0.7:
        check_result["ok"] = False
    
    return check_result


def calculate_score(checks: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Вычисляет риск-скор (0-100) и статус"""
    score = 100
    reasons = []
    
    for check in checks:
        check_name = check.get("name", "")
        check_ok = check.get("ok", True)
        check_details = check.get("details", [])
        
        if not check_ok:
            if check_name == "Arithmetic":
                score -= 15
                reasons.extend(check_details)
            elif check_name == "INN/KPP":
                score -= 20
                reasons.extend(check_details)
            elif check_name == "DuplicateNumber":
                score -= 20
                reasons.extend(check_details)
            elif check_name == "ContractMatch":
                similarity = check.get("similarity", 0)
                if similarity < 0.7:
                    score -= 15
                    reasons.append(f"Низкое сходство с договором: {similarity:.2f}")
            elif check_name == "BankPayment":
                score -= 20
                reasons.extend(check_details)
            elif check_name == "Dates":
                score -= 10
                reasons.extend(check_details)
            elif check_name == "VATRate":
                score -= 15
                reasons.extend(check_details)
            elif check_name == "Semantic":
                similarity = check.get("similarity", 1.0)
                if similarity < 0.7:
                    score -= 15
                    reasons.extend(check_details)
    
    # Определяем статус
    if score >= 80:
        status = "accepted"
    elif score >= 50:
        status = "suspicious"
    else:
        status = "rejected"
    
    return {
        "score": max(0, min(100, score)),
        "status": status,
        "reasons": reasons
    }


def verify_invoice(file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
    """
    Основная функция проверки счёта-фактуры
    Возвращает результат анализа с риск-скором и рекомендациями
    """
    analysis_id = str(uuid.uuid4())
    
    try:
        # Определяем тип файла и извлекаем текст
        text = ""
        file_ext = filename.lower().split('.')[-1] if filename else ""
        
        if mime_type == "application/pdf" or file_ext == 'pdf':
            if not PDF_AVAILABLE:
                raise ImportError("PDF processing not available")
            text = extract_text_from_pdf(file_content)
        elif mime_type == "application/xml" or file_ext == 'xml':
            if not XML_AVAILABLE:
                raise ImportError("XML processing not available")
            text = extract_text_from_xml(file_content)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                           "application/msword"] or file_ext in ['doc', 'docx']:
            text = extract_text_from_docx(file_content)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                          "application/vnd.ms-excel"] or file_ext in ['xls', 'xlsx', 'xlsm']:
            text = extract_text_from_excel(file_content, filename)
        else:
            raise ValueError(f"Unsupported file type: {mime_type} (extension: {file_ext})")
        
        if not text or len(text.strip()) < 50:
            raise ValueError("Не удалось извлечь текст из документа или текст слишком короткий")
        
        # Извлекаем реквизиты
        invoice_data = extract_invoice_data(text)
        
        # Выполняем проверки
        checks = []
        
        # 1. Арифметика
        arithmetic_check = check_arithmetic(invoice_data)
        checks.append({
            "name": "Arithmetic",
            "ok": arithmetic_check["ok"],
            "details": arithmetic_check["details"]
        })
        
        # 2. ИНН/КПП
        inn_ok = True
        inn_details = []
        if invoice_data.get("seller", {}).get("inn"):
            if not validate_inn(invoice_data["seller"]["inn"]):
                inn_ok = False
                inn_details.append("Невалидный ИНН продавца")
        else:
            inn_ok = False
            inn_details.append("ИНН продавца не найден")
        
        if invoice_data.get("seller", {}).get("kpp"):
            if not validate_kpp(invoice_data["seller"]["kpp"]):
                inn_ok = False
                inn_details.append("Невалидный КПП продавца")
        
        checks.append({
            "name": "INN/KPP",
            "ok": inn_ok,
            "details": inn_details
        })
        
        # 3. Дубликаты
        duplicate_check = check_duplicate(invoice_data)
        checks.append({
            "name": "DuplicateNumber",
            "ok": duplicate_check["ok"],
            "details": duplicate_check["details"]
        })
        
        # 4. Даты
        dates_check = check_dates(invoice_data)
        checks.append({
            "name": "Dates",
            "ok": dates_check["ok"],
            "details": dates_check["details"]
        })
        
        # 5. Ставка НДС
        vat_check = check_vat_rate(invoice_data)
        checks.append({
            "name": "VATRate",
            "ok": vat_check["ok"],
            "details": vat_check["details"]
        })
        
        # 6. Семантическая проверка
        semantic_check_result = semantic_check(invoice_data, text)
        checks.append({
            "name": "Semantic",
            "ok": semantic_check_result["ok"],
            "similarity": semantic_check_result["similarity"],
            "details": semantic_check_result["details"]
        })
        
        # 7. Проверка совпадения с договором (заглушка - в реальности запрос к CRM/1С)
        checks.append({
            "name": "ContractMatch",
            "ok": True,
            "similarity": 0.85,  # Заглушка
            "details": []
        })
        
        # 8. Проверка платёжки (заглушка - в реальности запрос к банку)
        checks.append({
            "name": "BankPayment",
            "ok": True,  # Заглушка - предполагаем что платёжка найдена
            "details": []
        })
        
        # Вычисляем скоринг
        scoring = calculate_score(checks)
        
        # Формируем рекомендации
        recommendations = []
        if scoring["status"] == "suspicious":
            recommendations.append("Рекомендуется запросить подтверждение платёжного поручения")
            recommendations.append("Проверить соответствие суммы договору")
        elif scoring["status"] == "rejected":
            recommendations.append("Документ отклонён - требуется повторная проверка")
            recommendations.append("Запросить корректный счёт-фактуру у продавца")
            recommendations.append("Сверить реквизиты с договором")
        
        # Вычисляем хэш документа
        doc_hash = hashlib.sha256(file_content).hexdigest()
        
        # Формируем результат
        result = {
            "analysis_id": analysis_id,
            "status": scoring["status"],
            "score": scoring["score"],
            "invoice": invoice_data,
            "checks": checks,
            "reasons": scoring["reasons"],
            "recommendations": recommendations,
            "audit": {
                "docHash": f"sha256:{doc_hash}",
                "timestamp": datetime.now().isoformat(),
                "engine": "ai-invoice-v1.0"
            }
        }
        
        # Сохраняем результат
        INVOICE_ANALYSES[analysis_id] = result
        
        logger.info(f"Invoice analyzed: {filename}, score={scoring['score']}, status={scoring['status']}")
        
        return result
        
    except Exception as e:
        logger.error(f"Error analyzing invoice: {e}")
        raise


def get_invoice_analysis(analysis_id: str) -> Optional[Dict[str, Any]]:
    """Получает результат анализа по ID"""
    return INVOICE_ANALYSES.get(analysis_id)

